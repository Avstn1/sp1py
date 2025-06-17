import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from openai import OpenAI
import os

try:
    from docx2pdf import convert
except ImportError:
    convert = None

api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    messagebox.showerror("API Key Error", "Please set your OPENAI_API_KEY environment variable.")
    exit()

client = OpenAI(api_key=api_key)

def get_skills_prompt(job_description):
    return f"""
You are writing a SKILLS section for a resume.

Formatting rules:
- Exactly 5 skills
- Each skill must be on a single line (no line breaks between label and description)
- No blank lines in between
- Each skill must start with a label followed by a colon (e.g. "Sales Skills:")
- The label and description must be on the same line e.g. (<Skill>: <Description>) (do not newline between label and description)
- Use plain text only (no bullets, bold, or italics)
- Each skill must reflect a real competency, not a job requirement
- Avoid phrases like “preferred” or “is a plus”
- Write in the tone of a resume, not a job posting

Tailor your responses to this job description:

{job_description}
"""

def get_experience_prompt(job_description, role_name, example_bullets):
    return f"""
You're writing bullet points for a resume's EXPERIENCE section. Each bullet must:
- Start with an action verb
- Stay authentic to the original experience
- Be concise
- Use • (the bullet symbol) at the beginning of each line
- Be limited to 3 bullet points UNLESS the job description clearly calls for a specific skill shown in the original experience. In that case, write 4.

Job Title: {role_name}

Original Experience:
{example_bullets}

Job Posting:
{job_description}

Now rewrite concise and relevant bullet points for this role.
"""

def ask_openai(prompt):
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        messagebox.showerror("OpenAI Error", str(e))
        return ""

def replace_skills_placeholder(doc, placeholder, skills_text):
    skills_lines = [line.strip() for line in skills_text.splitlines() if line.strip()]
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            parent = paragraph._element.getparent()
            p_idx = parent.index(paragraph._element)
            # original_style = paragraph.style  # Remove this
            parent.remove(paragraph._element)
            for i, skill_line in enumerate(skills_lines):
                new_p = doc.add_paragraph(skill_line)
                new_p.style = 'Normal'  # Force Normal style
                for run in new_p.runs:
                    run.font.size = Pt(10)
                p_format = new_p.paragraph_format
                p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(0)
                new_p_element = new_p._element
                parent.remove(new_p_element)
                parent.insert(p_idx + i, new_p_element)
            break

def replace_placeholders(doc_path, output_path, replacements):
    doc = Document(doc_path)
    for key, value in replacements.items():
        if key == "SKILLS":
            replace_skills_placeholder(doc, f"{{{{{key}}}}}", value)
        else:
            for paragraph in doc.paragraphs:
                if f"{{{{{key}}}}}" in paragraph.text:
                    for run in paragraph.runs:
                        if f"{{{{{key}}}}}" in run.text:
                            run.text = run.text.replace(f"{{{{{key}}}}}", value or "")
                    paragraph.style = 'Normal'  # force Normal style
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if f"{{{{{key}}}}}" in paragraph.text:
                                for run in paragraph.runs:
                                    if f"{{{{{key}}}}}" in run.text:
                                        run.text = run.text.replace(f"{{{{{key}}}}}", value or "")
                                paragraph.style = 'Normal'  # force Normal style
    doc.save(output_path)

def convert_docx_to_pdf(docx_path, pdf_path):
    if convert is None:
        messagebox.showwarning("Conversion Warning", "docx2pdf not installed. PDF conversion skipped.")
        return False
    try:
        convert(docx_path, pdf_path)
        return True
    except Exception as e:
        messagebox.showerror("PDF Conversion Error", f"Failed to convert DOCX to PDF:\n{str(e)}")
        return False

def generate():
    job_desc = job_text.get("1.0", tk.END).strip()
    pdf_name = pdf_entry.get().strip()

    if not job_desc:
        messagebox.showerror("Missing Info", "Please enter the job description.")
        return
    if not pdf_name:
        messagebox.showerror("Missing Filename", "Please enter a name for the output PDF.")
        return
    if not pdf_name.lower().endswith(".pdf"):
        pdf_name += ".pdf"

    generate_button.config(state=tk.DISABLED)
    root.update()

    skills_text = ask_openai(get_skills_prompt(job_desc))

    jdrf_example = "• Engaged customers in natural, friendly conversations to assess needs and provide exceptional sales service\n• Maintained personal and productivity goals to contribute to a positive work environment\n• Demonstrated expertise in products and trends to fit customer needs\n• Collaborated with team members to ensure high levels of customer satisfaction"
    doordash_example = "• Engaged customers in friendly and knowledgeable conversations to assess their needs\n• Exceeded personal and productivity goals to contribute to a positive sales environment\n• Maintained up-to-date product knowledge and trends to provide tailored recommendations\n• Collaborated with team members to ensure outstanding customer service and satisfaction"
    rev_example = "• Engaged with customers to deliver exceptional sales service and ensure high levels of satisfaction\n• Achieved personal and productivity goals while maintaining knowledge of all products offered\n• Adapted to different customer needs by asking open-ended questions and sharing expertise on products and trends\n• Contributed to a positive work environment by collaborating with team members and initiating tasks independently"
    camp_example = "• Engaged with customers to deliver an elevated shopping experience\n• Achieved personal and productivity goals while maintaining product knowledge\n• Collaborated with team members to provide excellent sales service\n• Adapted to different customer needs and resolved issues with a smile"

    exp_jdrf = ask_openai(get_experience_prompt(job_desc, "Software Engineer Co-op", jdrf_example))
    exp_door = ask_openai(get_experience_prompt(job_desc, "DoorDash Delivery Driver", doordash_example))
    exp_rev = ask_openai(get_experience_prompt(job_desc, "Rev Captionist", rev_example))
    exp_camp = ask_openai(get_experience_prompt(job_desc, "Church Children’s Camp Volunteer", camp_example))

    replacements = {
        "SKILLS": skills_text,
        "JDRF": exp_jdrf,
        "DOORDASH": exp_door,
        "REV": exp_rev,
        "CAMP": exp_camp
    }

    try:
        replace_placeholders("resume_template.docx", "output_resume.docx", replacements)
        pdf_success = convert_docx_to_pdf("output_resume.docx", pdf_name)

        msg = f"Resume generated as 'output_resume.docx'."
        if pdf_success:
            msg += f"\nAlso saved as '{pdf_name}'."
        else:
            msg += "\nPDF conversion was skipped or failed."

        messagebox.showinfo("Success", msg)
    except Exception as e:
        messagebox.showerror("Error", f"Failed to generate resume: {str(e)}")
    finally:
        generate_button.config(state=tk.NORMAL)

# ---------------- UI Styling -------------------
root = tk.Tk()
root.title("Resume Generator")
root.geometry("950x550")
root.configure(bg="#1e1e1e")

font_main = ("Segoe UI", 10)

style = ttk.Style(root)
style.theme_use("clam")

style.configure("TFrame", background="#1e1e1e")
style.configure("TLabel", background="#1e1e1e", foreground="#dcdcdc", font=("Segoe UI", 11))
style.configure("TEntry", fieldbackground="#ffffff", foreground="#000000", font=("Segoe UI", 10))
style.configure("TButton", background="#444", foreground="white", font=("Segoe UI", 10), padding=6)
style.map("TButton", background=[("active", "#5a5a5a")])

main_frame = ttk.Frame(root, padding="15", style="TFrame")
main_frame.pack(fill="both", expand=True)

ttk.Label(main_frame, text="Job Description:").grid(row=0, column=0, sticky="nw", pady=(0, 5))
job_text = tk.Text(main_frame, width=85, height=15, bg="#ffffff", fg="#000000", insertbackground="black", font=font_main, bd=1, relief="solid")
job_text.grid(row=1, column=0, columnspan=2, pady=(0, 10))

ttk.Label(main_frame, text="PDF File Name:").grid(row=2, column=0, sticky="w")
pdf_entry = ttk.Entry(main_frame, width=50)
pdf_entry.insert(0, "Austin Bartolome - Resume.pdf")
pdf_entry.grid(row=2, column=1, sticky="w", pady=5)

generate_button = ttk.Button(main_frame, text="Generate Resume", command=generate)
generate_button.grid(row=3, column=1, sticky="e", pady=15)

main_frame.columnconfigure(0, weight=1)
main_frame.columnconfigure(1, weight=1)

root.mainloop()
